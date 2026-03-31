#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
每日阅读更新脚本
用法: python3 update.py <word文件路径> [标题]
"""

import os
import sys
import json
import re
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_table_html(table):
    """将 Word 表格转换为 HTML 字符串"""
    rows = []
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # 提取单元格内所有段落文本，用 <br> 连接
            cell_text = '<br>'.join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            cells.append(cell_text)
        if i == 0:
            rows.append('<tr>' + ''.join(f'<th>{c}</th>' for c in cells) + '</tr>')
        else:
            rows.append('<tr>' + ''.join(f'<td>{c}</td>' for c in cells) + '</tr>')
    return '<table>' + '\n'.join(rows) + '</table>'


def extract_content(doc):
    """
    按文档顺序提取段落和表格，返回混合 HTML 内容。
    - 段落 → 根据样式判断层级
    - 表格 → <table>...</table>
    - 引用块 → 用 blockquote 包裹
    - 代码块 → 用 <pre> 包裹
    """
    parts = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # 找到对应的 Paragraph 对象
            from docx.oxml.ns import qn
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ''

            # Heading 1 / Heading 2 / Heading 3
            if 'Heading 1' in style_name or style_name == 'heading 1':
                parts.append(f'<h2>{text}</h2>')
            elif 'Heading 2' in style_name or style_name == 'heading 2':
                parts.append(f'<h3>{text}</h3>')
            elif 'Heading 3' in style_name or style_name == 'heading 3':
                parts.append(f'<h4>{text}</h4>')
            elif 'Heading' in style_name:
                parts.append(f'<h3>{text}</h3>')
            # Quote style or quote-like
            elif 'Quote' in style_name or 'Block' in style_name:
                parts.append(f'<blockquote>{text}</blockquote>')
            elif 'Title' in style_name:
                parts.append(f'<h1>{text}</h1>')
            else:
                parts.append(f'<p>{text}</p>')

        elif tag == 'tbl':
            # 找到对应的 Table 对象
            from docx.table import Table
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table is not None:
                parts.append(extract_table_html(table))

    return '\n'.join(parts)


def update_article(word_path, title=None):
    """更新文章到网站"""

    if not os.path.exists(word_path):
        print(f'错误: 文件不存在 - {word_path}')
        sys.exit(1)

    # 读取Word文件
    doc = Document(word_path)

    # 提取混合内容（段落 + 表格 HTML）
    content = extract_content(doc)

    # 如果提取不到内容，回退到纯文本模式
    if not content.strip():
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(f'<p>{text}</p>')
        content = '\n'.join(paragraphs)

    # 生成日期和文件名
    today = datetime.now().strftime('%Y-%m-%d')
    json_filename = f'{today}.json'

    # 如果没有提供标题，尝试从内容提取
    if not title:
        # 尝试找到第一个 h1 或第一段文本
        h1_match = re.search(r'<h1>(.*?)</h1>', content)
        if h1_match:
            title = h1_match.group(1)
        else:
            first_text = re.search(r'<p>(.*?)</p>', content)
            title = first_text.group(1) if first_text else today

    # 文章JSON内容
    article_data = {
        'date': today,
        'title': title,
        'content': content
    }

    # 保存文章JSON文件
    articles_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'articles')
    os.makedirs(articles_dir, exist_ok=True)
    json_path = os.path.join(articles_dir, json_filename)

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(article_data, f, ensure_ascii=False, indent=2)

    print(f'已保存文章: {json_path}')

    # 更新文章列表
    list_path = os.path.join(articles_dir, 'list.json')

    if os.path.exists(list_path):
        with open(list_path, 'r', encoding='utf-8') as f:
            articles_list = json.load(f)
    else:
        articles_list = []

    # 检查今天是否已有文章，有则更新，无则添加
    existing_index = None
    for i, article in enumerate(articles_list):
        if article.get('date') == today:
            existing_index = i
            break

    article_info = {
        'date': today,
        'title': title,
        'file': json_filename
    }

    if existing_index is not None:
        articles_list[existing_index] = article_info
        print(f'已更新今天的文章')
    else:
        articles_list.append(article_info)
        print(f'已添加新文章')

    # 按日期排序（最新在前）
    articles_list.sort(key=lambda x: x.get('date', ''), reverse=True)

    with open(list_path, 'w', encoding='utf-8') as f:
        json.dump(articles_list, f, ensure_ascii=False, indent=2)

    print(f'已更新文章列表，共 {len(articles_list)} 篇文章')
    print(f'\n请运行以下命令部署到GitHub:')
    print(f'  cd {os.path.dirname(os.path.abspath(__file__))}')
    print(f'  git add .')
    print(f'  git commit -m "更新文章 {today}"')
    print(f'  git push')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python3 update.py <word文件路径> [标题]')
        sys.exit(1)

    word_path = sys.argv[1]
    title = sys.argv[2] if len(sys.argv) > 2 else None

    update_article(word_path, title)